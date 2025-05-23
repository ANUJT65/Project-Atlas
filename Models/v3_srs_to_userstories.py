import io
import base64
import json
from pathlib import Path
import re
from openai import AzureOpenAI
import PyPDF2
import pandas as pd
import docx
from PIL import Image
import requests
import os
from azure.ai.formrecognizer import DocumentAnalysisClient
from azure.core.credentials import AzureKeyCredential
import httpx
import tempfile
from azure.cosmos import CosmosClient, exceptions
from datetime import datetime

# Azure OpenAI Configuration
AZURE_ENDPOINT = "https://suraj-m9lgdbv9-eastus2.cognitiveservices.azure.com/"
AZURE_API_KEY = "75PVa3SAy9S2ZR590gZesyTNDMZtb3Oa5EdHlRbqWeQ89bmoOGl4JQQJ99BDACHYHv6XJ3w3AAAAACOGUP8d"
AZURE_API_VERSION = "2024-02-15-preview"
DEPLOYMENT_NAME = "gpt-4v"  # Make sure this is your vision model deployment name
TEXT_DEPLOYMENT_NAME = "gpt-4"  # Your text model deployment name

# Azure Document Intelligence Configuration
FORM_ENDPOINT = "https://barclaysform.cognitiveservices.azure.com/"
FORM_KEY = "63spGg0VYFV0kWZB3nmsFDp8yEbi40zmEnCvIl6D8Seih4YyLsp9JQQJ99BDACYeBjFXJ3w3AAALACOGh5hu"

# Azure Cosmos DB Configuration
COSMOS_ENDPOINT = "https://barclaysdb.documents.azure.com:443/"
COSMOS_KEY = "ercuc7wFNt4RPsxcx2QTzKP8AhKUDnjJrt0mpZrW2Yi2IQGAa7wDrEbhKRHsurGu0P7GuGny4IZRACDbtecfNQ=="
COSMOS_DATABASE = "RequirementsDB"  # Changed to match existing database
COSMOS_CONTAINER = "UserStories"    # Changed to match existing container

# Initialize clients
openai_client = AzureOpenAI(
    api_version=AZURE_API_VERSION,
    api_key=AZURE_API_KEY,
    azure_endpoint=AZURE_ENDPOINT
)

document_client = DocumentAnalysisClient(
    endpoint=FORM_ENDPOINT, 
    credential=AzureKeyCredential(FORM_KEY)
)

# Initialize Cosmos DB client with error handling
try:
    cosmos_client = CosmosClient(COSMOS_ENDPOINT, COSMOS_KEY)
    database = cosmos_client.get_database_client(COSMOS_DATABASE)
    container = database.get_container_client(COSMOS_CONTAINER)
    print("Successfully connected to Cosmos DB")
except Exception as e:
    print(f"Error connecting to Cosmos DB: {e}")
    cosmos_client = None
    database = None
    container = None

def is_url(path_or_url: str) -> bool:
    """
    Check if the input is a URL.
    """
    return path_or_url.startswith("http://") or path_or_url.startswith("https://")

def download_from_url(url: str) -> Path:
    """
    Download a file from a URL and return the local path.
    """
    try:
        response = httpx.get(url)
        response.raise_for_status()
        
        # Create a temporary file with the appropriate extension
        suffix = Path(url).suffix
        if not suffix:
            # Default to .pdf if no extension is found
            suffix = ".pdf"
            
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
        temp_file.write(response.content)
        temp_file.close()
        
        return Path(temp_file.name)
    except Exception as e:
        print(f"Error downloading file from URL: {e}")
        raise

def extract_with_document_intelligence(path: Path) -> str:
    """
    Extract text from any document using Azure Document Intelligence.
    This provides better extraction for complex documents with tables, forms, etc.
    """
    try:
        with open(path, "rb") as document:
            poller = document_client.begin_analyze_document(
                "prebuilt-document", document
            )
            result = poller.result()

            # Extract all text content
            text_content = []
            
            # Get all paragraphs
            for paragraph in result.paragraphs:
                text_content.append(paragraph.content)
            
            # Get text from tables
            for table in result.tables:
                table_text = []
                for cell in table.cells:
                    table_text.append(cell.content)
                text_content.append(" | ".join(table_text))
            
            # Get key-value pairs
            for kv_pair in result.key_value_pairs:
                if kv_pair.key and kv_pair.value:
                    text_content.append(f"{kv_pair.key.content}: {kv_pair.value.content}")

            return "\n".join(text_content)
    except Exception as e:
        print(f"Error in Document Intelligence processing: {e}")
        return ""

def extract_from_url(url: str) -> str:
    """
    Extract text from a document URL using Azure Document Intelligence.
    """
    try:
        # Download the file from the URL
        local_path = download_from_url(url)
        
        # Extract text using Document Intelligence
        text = extract_with_document_intelligence(local_path)
        
        # Clean up the temporary file
        try:
            os.unlink(local_path)
        except:
            pass
            
        return text
    except Exception as e:
        print(f"Error extracting text from URL: {e}")
        return ""

def load_file_text(file_path: str) -> str:
    """
    Extract text from various file types or URLs.
    First attempts Azure Document Intelligence, then falls back to specific extractors.
    """
    # Check if the input is a URL
    if is_url(file_path):
        print(f"Processing document from URL: {file_path}")
        return extract_from_url(file_path)
    
    # Process local file
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"{file_path} does not exist.")
    
    suffix = path.suffix.lower()
    
    try:
        # First try Azure Document Intelligence for supported formats
        if suffix in ['.pdf', '.docx', '.doc', '.jpeg', '.jpg', '.png', '.bmp', '.tiff', '.tif']:
            print("Using Azure Document Intelligence for extraction...")
            text = extract_with_document_intelligence(path)
            if text.strip():
                return text
            print("Falling back to specific extractors...")
        
        # Fall back to specific extractors
        if suffix in ['.jpg', '.jpeg', '.png', '.gif', '.bmp']:
            return process_image(path)
        elif suffix == '.pdf':
            return extract_pdf_text(path)
        elif suffix == '.docx':
            return extract_docx_text(path)
        elif suffix in ['.xls', '.xlsx']:
            return extract_excel_text(path)
        elif suffix == '.txt':
            return path.read_text(encoding="utf-8")
        else:
            raise ValueError(f"Unsupported file type: {suffix}")
    except Exception as e:
        print(f"Error processing {file_path}: {e}")
        return ""

def process_image(path: Path) -> str:
    """
    Process an image file: resize if needed, encode to base64,
    and have the model describe it.
    """
    with Image.open(path) as img:
        max_size = (800, 800)
        img.thumbnail(max_size)
        buffer = io.BytesIO()
        fmt = img.format if img.format else "JPEG"
        img.save(buffer, format=fmt)
        buffer.seek(0)
    
    image_data = base64.b64encode(buffer.getvalue()).decode('utf-8')
    
    try:
        response = openai_client.chat.completions.create(
            model=DEPLOYMENT_NAME,
            messages=[
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": "Thoroughly describe the contents of this image. Extract any visible text and provide a comprehensive analysis."},
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:image/jpeg;base64,{image_data}"
                            }
                        }
                    ]
                }
            ],
            max_tokens=1000
        )
        return response.choices[0].message.content
    except Exception as e:
        print(f"Error processing image: {e}")
        return ""

def extract_pdf_text(path: Path) -> str:
    """Extract text from a PDF file."""
    with path.open("rb") as file:
        reader = PyPDF2.PdfReader(file)
        texts = [page.extract_text() for page in reader.pages if page.extract_text()]
        return "\n".join(texts)

def extract_docx_text(path: Path) -> str:
    """Extract text from a DOCX file."""
    doc = docx.Document(path)
    return "\n".join(paragraph.text for paragraph in doc.paragraphs)

def extract_excel_text(path: Path) -> str:
    """Extract text from an Excel file."""
    df = pd.read_excel(path)
    return df.to_string()

def extract_user_stories_json(document: str) -> list:
    """
    Generate functional user stories from a technical document using Azure OpenAI.
    Returns a list of user story dictionaries.
    """
    prompt = (
        "Extract user stories from the following technical requirement document and output ONLY a valid JSON array. "
        "Each element in the array must be a JSON object with the keys: 'title', 'user_role', 'goal', 'benefit', "
        "'acceptance_criteria', 'priority', and 'document'. "
        "The 'priority' field should be one of: 'Must Have', 'Should Have', 'Could Have', or 'Won't Have'. "
        "Do not output anything else besides the JSON array.\n\n"
        "Document:\n\n{document}\n\n"
        "Guidelines:\n"
        "- Extract only functional user stories.\n"
        "- Ensure clarity and avoid technical jargon.\n"
        "- Group related stories if applicable.\n"
        "- Use the MoSCoW prioritization framework for the priority field."
    ).format(document=document)
    
    try:
        response = openai_client.chat.completions.create(
            model="gpt-4o",  # Use your text model deployment
            messages=[
                {"role": "system", "content": "You are a requirements analyst that creates user stories from technical documents. Always respond with valid JSON only."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.3,
            max_tokens=3000
        )
        
        cleaned_text = response.choices[0].message.content.replace("```json", "").replace("```", "").strip()
        
        # Attempt to parse the JSON
        try:
            stories = json.loads(cleaned_text)
            return stories
        except json.JSONDecodeError:
            # Attempt to extract the JSON array substring
            match = re.search(r'\[.*\]', cleaned_text, re.DOTALL)
            if match:
                json_text = match.group(0)
                try:
                    stories = json.loads(json_text)
                    return stories
                except json.JSONDecodeError as e:
                    print("Warning: Failed to parse JSON after substring extraction:", e)
                    print("Extracted substring:")
                    print(json_text)
                    # As a final hack, if missing a closing bracket, add it.
                    if not json_text.strip().endswith("]"):
                        fixed_text = json_text.strip() + "]"
                        try:
                            stories = json.loads(fixed_text)
                            return stories
                        except Exception as e:
                            print("Warning: Failed to fix JSON by appending closing bracket:", e)
                    return []
            else:
                print("Warning: No JSON array found in the response.")
                print("Raw cleaned response:")
                print(cleaned_text)
                return []
    except Exception as e:
        print(f"Error generating user stories: {e}")
        return []

def push_story_to_jira(story: dict):
    """
    Push a user story to Jira using direct credentials.
    """
    # Direct JIRA credentials
    jira_domain = "prathamgadkari.atlassian.net"
    jira_email = "prathamgadkari@gmail.com"
    jira_api_token = "ATATT3xFfGF0LXSWvbAlLm1lVXyxSroefrBXONpTLDg5mzfknXUUxTYAjQ0u59wKKGh3ObCEnducPVGqRwCLoxDu8oc4xx3aBAHj6A9Tzuj2OqiUszHVo3UvYrpmblYC2xHttFxo5ULieeRE2LNIfR3w_l2YNJTvACQ_zmBnkt6Tjvenqyu2pEM=99E9EA51"
    project_key =  "SCRUM"

    jira_url = f"https://{jira_domain}/rest/api/2/issue"
    auth = (jira_email, jira_api_token)
    headers = {"Content-Type": "application/json"}
    
    payload = {
        "fields": {
            "project": {"key": project_key},
            "summary": story.get("title", "User Story"),
            "description": json.dumps(story, indent=2),
            "issuetype": {"name": "Story"}
        }
    }
    try:
        response = requests.post(jira_url, auth=auth, headers=headers, data=json.dumps(payload))
        if response.status_code == 201:
            print(f"Story '{story.get('title')}' successfully pushed to Jira.")
        else:
            print(f"Failed to push story '{story.get('title')}' to Jira. Status code: {response.status_code}. Response: {response.text}")
    except Exception as e:
        print(f"Error pushing story '{story.get('title')}' to Jira: {e}")

def store_story_in_cosmos(story: dict, project_id: str, standard_doc_id: str) -> str:
    """
    Store a user story in Cosmos DB following the specified structure.
    Returns the story ID.
    """
    if not container:
        print("Cosmos DB connection not available. Skipping story storage.")
        return None
        
    try:
        # Generate a unique story ID
        story_id = f"story_{project_id}_{standard_doc_id}_{int(datetime.now().timestamp())}"
        
        # Create the story document with all fields from the JSON structure
        story_doc = {
            "id": story_id,
            "project_id": project_id,
            "standard_doc_id": standard_doc_id,
            "title": story.get("title", "Untitled Story"),
            "user_role": story.get("user_role", ""),
            "goal": story.get("goal", ""),
            "benefit": story.get("benefit", ""),
            "acceptance_criteria": story.get("acceptance_criteria", ""),
            "priority": story.get("priority", "Must Have"),
            "document": story.get("document", ""),
            "jira_issue_id": story.get("jira_issue_id", ""),
            "status": "To Do",
            "created_at": datetime.utcnow().isoformat(),
        }
        
        # Create the item in Cosmos DB
        container.create_item(body=story_doc)
        print(f"Story '{story_doc['title']}' successfully stored in Cosmos DB.")
        return story_id
        
    except exceptions.CosmosHttpResponseError as e:
        print(f"Cosmos DB error storing story: {e}")
        return None
    except Exception as e:
        print(f"Error storing story in Cosmos DB: {e}")
        return None

def main():
    file_path = input("Enter the path to your document or URL: ").strip()
    project_id = input("Enter the project ID: ").strip()
    standard_doc_id = input("Enter the standard document ID: ").strip()
    
    try:
        document_text = load_file_text(file_path)
    except Exception as e:
        print(f"Failed to load document: {e}")
        return
    
    if not document_text:
        print("No text could be extracted from the document.")
        return

    print("\nGenerating user stories in JSON format from the document...\n")
    stories = extract_user_stories_json(document_text)
    if stories:
        print("Extracted User Stories (JSON):\n")
        print(json.dumps(stories, indent=2))
        
        # Save JSON output for reference
        output_file = Path("user_stories.json")
        with output_file.open("w", encoding="utf-8") as f:
            json.dump(stories, f, indent=2)
        print(f"\nUser stories saved to '{output_file}'.")
        
        # Store stories in Cosmos DB if connection is available
        if container:
            print("\nStoring stories in Cosmos DB...")
            for story in stories:
                story_id = store_story_in_cosmos(story, project_id, standard_doc_id)
                if story_id:
                    story["id"] = story_id
        else:
            print("\nSkipping Cosmos DB storage due to connection issues.")
        
        # Wait for user input to push to Jira
        push_input = input("\nPress 'y' to push these stories to Jira, or any other key to exit: ").strip().lower()
        if push_input == "y":
            for story in stories:
                push_story_to_jira(story)
        else:
            print("Skipping Jira push.")
    else:
        print("No user stories could be extracted.")

if __name__ == "__main__":
    main()